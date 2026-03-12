"""
DOCX Generator MCP Server
python-docx 기반 마크다운 -> Word 문서 변환 서버
샘플 기준: 파란 섹션 헤더, 회색 코드블록, 깔끔한 비즈니스 문서 스타일
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any

# MCP SDK
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# DOCX 생성
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# 출력 디렉토리
OUTPUT_DIR = Path(__file__).parent.parent.parent.parent / "data" / "docx_output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# MCP 서버 생성
server = Server("docx-generator")

# ============================================================================
# 스타일 정의
# ============================================================================

STYLES = {
    "technical": {
        "title_color": RGBColor(0x1A, 0x36, 0x5D),       # 진한 네이비
        "h2_color": RGBColor(0x2B, 0x6C, 0xB0),          # 파란색
        "h3_color": RGBColor(0x2B, 0x6C, 0xB0),          # 파란색
        "text_color": RGBColor(0x33, 0x33, 0x33),         # 진한 회색
        "table_header_bg": "2B6CB0",                       # 파란색
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),  # 흰색
        "table_row_even_bg": "F7FAFC",                     # 연한 파란 회색
        "table_border_color": "B0C4DE",                    # 연한 파란 테두리
        "code_bg": "F0F4F8",                               # 연한 회색 배경
        "code_text": RGBColor(0x2D, 0x3A, 0x4A),          # 진한 회색
        "blockquote_bg": "FFF8E1",                         # 연한 노란색
        "blockquote_border": "F0C040",                     # 노란색
        "warning_color": RGBColor(0xC0, 0x60, 0x00),      # 주황색 (경고)
    },
    "report": {
        "title_color": RGBColor(0x1A, 0x1A, 0x1A),
        "h2_color": RGBColor(0x33, 0x33, 0x33),
        "h3_color": RGBColor(0x44, 0x44, 0x44),
        "text_color": RGBColor(0x22, 0x22, 0x22),
        "table_header_bg": "4A5568",
        "table_header_text": RGBColor(0xFF, 0xFF, 0xFF),
        "table_row_even_bg": "F7FAFC",
        "table_border_color": "CBD5E0",
        "code_bg": "F7FAFC",
        "code_text": RGBColor(0x2D, 0x3A, 0x4A),
        "blockquote_bg": "F7FAFC",
        "blockquote_border": "A0AEC0",
        "warning_color": RGBColor(0xC0, 0x60, 0x00),
    },
    "simple": {
        "title_color": RGBColor(0x33, 0x33, 0x33),
        "h2_color": RGBColor(0x44, 0x44, 0x44),
        "h3_color": RGBColor(0x55, 0x55, 0x55),
        "text_color": RGBColor(0x33, 0x33, 0x33),
        "table_header_bg": "E2E8F0",
        "table_header_text": RGBColor(0x2D, 0x3A, 0x4A),
        "table_row_even_bg": "F7FAFC",
        "table_border_color": "E2E8F0",
        "code_bg": "F5F5F5",
        "code_text": RGBColor(0x33, 0x33, 0x33),
        "blockquote_bg": "F5F5F5",
        "blockquote_border": "CCCCCC",
        "warning_color": RGBColor(0xC0, 0x60, 0x00),
    },
}

FONT_NAME = "맑은 고딕"
CODE_FONT_NAME = "Consolas"


# ============================================================================
# 마크다운 파서 (PDF 서버와 동일 로직)
# ============================================================================

def strip_markdown_formatting(text: str) -> str:
    """인라인 마크다운 포맷팅 제거"""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def _is_list_line(line: str) -> bool:
    stripped = line.strip()
    return bool(re.match(r'^[-*+]\s+', stripped) or re.match(r'^\d+\.\s+', stripped))


def _has_consecutive_numbered(lines: List[str], idx: int) -> bool:
    for j in range(idx + 1, len(lines)):
        stripped = lines[j].strip()
        if not stripped:
            continue
        return bool(re.match(r'^\d+\.\s+', stripped))
    return False


def _collect_list_items(lines: List[str], start_idx: int) -> tuple:
    items = []
    i = start_idx
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and _is_list_line(lines[j]):
                i = j
                continue
            break
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if ul_match:
            indent = len(ul_match.group(1))
            depth = indent // 2
            items.append({'text': ul_match.group(2).strip(), 'depth': depth, 'ordered': False})
            i += 1
            continue
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol_match:
            indent = len(ol_match.group(1))
            depth = indent // 2
            items.append({'text': ol_match.group(2).strip(), 'depth': depth, 'ordered': True})
            i += 1
            continue
        break
    return items, i


def parse_markdown_content(content: str) -> List[Dict[str, Any]]:
    """마크다운 컨텐츠를 구조화된 요소로 파싱"""
    elements = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # 이미지
        img_match = re.match(r'!\[([^\]]*)\]\(([^)\s]+)(?:\s+"([^"]*)")?\)', stripped)
        if img_match:
            elements.append({
                'type': 'image',
                'path': img_match.group(2),
                'alt': img_match.group(1),
                'caption': img_match.group(3) or img_match.group(1),
            })
            i += 1
            continue

        # 제목
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped.lstrip('#').strip()
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # 코드 블록
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append({'type': 'code', 'lang': lang, 'text': '\n'.join(code_lines)})
            i += 1
            continue

        # 테이블
        if '|' in stripped and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [stripped]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            elements.append({'type': 'table', 'lines': table_lines})
            continue

        # 구분선
        if stripped in ['---', '***', '___']:
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 블록인용
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if s.startswith('>'):
                    text = re.sub(r'^>\s?', '', s)
                    quote_lines.append(text)
                    i += 1
                elif not s:
                    if i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                        quote_lines.append('')
                        i += 1
                    else:
                        break
                else:
                    break
            elements.append({'type': 'blockquote', 'text': '\n'.join(quote_lines)})
            continue

        # 불릿 리스트
        bullet_match = re.match(r'^(\s*)[-*+]\s+', line)
        if bullet_match:
            items, i = _collect_list_items(lines, i)
            if items:
                elements.append({'type': 'list', 'ordered': False, 'items': items})
            continue

        # 서브 섹션 (1-1. 2-3-1. 등)
        sub_section_match = re.match(r'^(\d+-\d+(?:-\d+)*)\.\s+(.+)$', stripped)
        if sub_section_match:
            section_num = sub_section_match.group(1)
            text = stripped
            level = 3 if section_num.count('-') == 1 else 4
            elements.append({'type': 'heading', 'level': level, 'text': text})
            i += 1
            continue

        # 번호 패턴
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            if _has_consecutive_numbered(lines, i):
                items, i = _collect_list_items(lines, i)
                if items:
                    elements.append({'type': 'list', 'ordered': True, 'items': items})
                continue
            else:
                elements.append({'type': 'heading', 'level': 2, 'text': stripped})
                i += 1
                continue

        # 일반 텍스트
        text_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (not next_line or
                next_line.startswith('#') or
                next_line.startswith('```') or
                next_line.startswith('>') or
                next_line in ['---', '***', '___'] or
                re.match(r'^[-*+]\s+', next_line) or
                re.match(r'^\d+(?:-\d+)*\.\s+', next_line) or
                re.match(r'^!\[', next_line) or
                ('|' in next_line and i + 1 < len(lines) and '---' in lines[i + 1])):
                break
            text_lines.append(next_line)
            i += 1
        combined_text = ' '.join(text_lines)
        elements.append({'type': 'paragraph', 'text': combined_text})

    return elements


def parse_table(table_lines: List[str]) -> Dict[str, Any]:
    """테이블 파싱"""
    rows = []
    for line in table_lines:
        if '---' in line and '|' in line:
            continue
        cells = [cell.strip() for cell in line.split('|')]
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        if cells:
            rows.append(cells)
    if len(rows) >= 1:
        return {'headers': rows[0], 'data': rows[1:] if len(rows) > 1 else []}
    return {'headers': [], 'data': []}


# ============================================================================
# DOCX 렌더링
# ============================================================================

def _set_cell_shading(cell, color_hex: str):
    """테이블 셀 배경색 설정"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, color_hex: str = "B0C4DE"):
    """테이블 셀 테두리 설정"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(borders)


def _add_inline_formatted_text(paragraph, text: str, style_colors: dict, base_size: int = 10, is_code_block: bool = False):
    """인라인 마크다운 서식을 적용하여 텍스트 추가 (**bold**, *italic*, `code`)"""
    # 패턴: **bold**, *italic*, `code`, 일반 텍스트
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))'
    font_name = CODE_FONT_NAME if is_code_block else FONT_NAME

    for match in re.finditer(pattern, text):
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(base_size)
            run.font.color.rgb = style_colors["text_color"]
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.name = font_name
            run.font.size = Pt(base_size)
            run.font.color.rgb = style_colors["text_color"]
        elif match.group(4):  # `code`
            run = paragraph.add_run(match.group(4))
            run.font.name = CODE_FONT_NAME
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = style_colors.get("code_text", style_colors["text_color"])
            # 인라인 코드 배경 (shading)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{style_colors["code_bg"]}" w:val="clear"/>')
            run._r.get_or_add_rPr().append(shading)
        elif match.group(5):  # 일반 텍스트
            run = paragraph.add_run(match.group(5))
            run.font.name = font_name
            run.font.size = Pt(base_size)
            if is_code_block:
                run.font.color.rgb = style_colors.get("code_text", style_colors["text_color"])
            else:
                run.font.color.rgb = style_colors["text_color"]


def _set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line_spacing: float = 1.15):
    """단락 간격 설정"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_horizontal_rule(doc, style_colors: dict):
    """수평선 추가"""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=6)
    # 하단 테두리로 수평선 표현
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBorders>'
    )
    pPr.append(borders)


def render_docx(
    elements: List[Dict[str, Any]],
    title: str,
    subtitle: str = "",
    style: str = "technical",
    section_per_page: bool = False,
) -> Document:
    """구조화된 요소를 DOCX로 렌더링"""
    doc = Document()
    colors = STYLES.get(style, STYLES["technical"])

    # ---- 기본 폰트 설정 ----
    doc_style = doc.styles['Normal']
    doc_style.font.name = FONT_NAME
    doc_style.font.size = Pt(10)
    doc_style.font.color.rgb = colors["text_color"]
    doc_style.paragraph_format.line_spacing = 1.15

    # 한글 폰트 설정 (eastAsia)
    doc_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # ---- 페이지 여백 설정 ----
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- 요소 렌더링 ----
    h1_rendered = False
    first_h2_after_h1 = True
    ordered_counters: Dict[int, int] = {}

    for elem in elements:
        elem_type = elem['type']

        if elem_type == 'heading':
            level = elem['level']
            text = elem['text']

            if level == 1:
                # 대제목: 중앙 정렬, 큰 볼드
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_spacing(p, before=20, after=4)
                run = p.add_run(strip_markdown_formatting(text))
                run.bold = True
                run.font.size = Pt(22)
                run.font.name = FONT_NAME
                run.font.color.rgb = colors["title_color"]

                # 부제목
                if subtitle:
                    p_sub = doc.add_paragraph()
                    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _set_paragraph_spacing(p_sub, before=0, after=12)
                    run_sub = p_sub.add_run(subtitle)
                    run_sub.italic = True
                    run_sub.font.size = Pt(11)
                    run_sub.font.name = FONT_NAME
                    run_sub.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                # 구분선
                _add_horizontal_rule(doc, colors)
                h1_rendered = True

            elif level == 2:
                # 섹션별 페이지 나누기
                if section_per_page and h1_rendered:
                    if first_h2_after_h1:
                        first_h2_after_h1 = False
                    else:
                        doc.add_page_break()

                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=16, after=6)
                run = p.add_run(strip_markdown_formatting(text))
                run.bold = True
                run.font.size = Pt(15)
                run.font.name = FONT_NAME
                run.font.color.rgb = colors["h2_color"]

                # h2 밑줄 (하단 테두리)
                pPr = p._p.get_or_add_pPr()
                h2_color = colors["h2_color"]
                hex_color = f"{h2_color[0]:02X}{h2_color[1]:02X}{h2_color[2]:02X}" if isinstance(h2_color, tuple) else f"{h2_color.red:02X}{h2_color.green:02X}{h2_color.blue:02X}"
                borders = parse_xml(
                    f'<w:pBorders {nsdecls("w")}>'
                    f'  <w:bottom w:val="single" w:sz="8" w:space="2" w:color="{hex_color}"/>'
                    f'</w:pBorders>'
                )
                pPr.append(borders)

            elif level == 3:
                # 서브섹션 제목
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=10, after=4)
                run = p.add_run(strip_markdown_formatting(text))
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = FONT_NAME
                run.font.color.rgb = colors["h3_color"]

            else:
                # h4+
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=8, after=3)
                run = p.add_run(strip_markdown_formatting(text))
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT_NAME
                run.font.color.rgb = colors["text_color"]

        elif elem_type == 'paragraph':
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=4)
            _add_inline_formatted_text(p, elem['text'], colors, base_size=10)

        elif elem_type == 'list':
            items = elem['items']
            ordered = elem['ordered']
            ordered_counters.clear()

            for item in items:
                depth = item.get('depth', 0)
                is_ordered = item.get('ordered', ordered)
                text = item['text']

                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=1)

                # 들여쓰기
                indent_cm = 1.0 + depth * 0.8
                p.paragraph_format.left_indent = Cm(indent_cm)
                p.paragraph_format.first_line_indent = Cm(-0.5)

                # 프리픽스
                if is_ordered:
                    ordered_counters.setdefault(depth, 0)
                    ordered_counters[depth] += 1
                    for d in list(ordered_counters.keys()):
                        if d > depth:
                            del ordered_counters[d]
                    prefix = f"{ordered_counters[depth]}.  "
                else:
                    bullets = ['\u2022', '\u2013', '\u00B7']  # bullet, en-dash, middle dot
                    bullet_char = bullets[min(depth, len(bullets) - 1)]
                    prefix = f"{bullet_char}  "

                run_prefix = p.add_run(prefix)
                run_prefix.font.name = FONT_NAME
                run_prefix.font.size = Pt(10)
                run_prefix.font.color.rgb = colors["text_color"]

                _add_inline_formatted_text(p, text, colors, base_size=10)

        elif elem_type == 'table':
            table_data = parse_table(elem['lines'])
            if not table_data['headers']:
                continue

            headers = table_data['headers']
            data = table_data['data']
            num_cols = len(headers)

            table = doc.add_table(rows=1 + len(data), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 테이블 전체 너비 설정
            table.autofit = True

            border_color = colors["table_border_color"]

            # 헤더 행
            for col_idx, header_text in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_spacing(p, before=2, after=2)
                run = p.add_run(strip_markdown_formatting(header_text))
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = FONT_NAME
                run.font.color.rgb = colors["table_header_text"]
                _set_cell_shading(cell, colors["table_header_bg"])
                _set_cell_border(cell, border_color)

            # 데이터 행
            for row_idx, row_data in enumerate(data):
                for col_idx in range(num_cols):
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                    cell.text = ""
                    p = cell.paragraphs[0]
                    align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    p.alignment = align
                    _set_paragraph_spacing(p, before=1, after=1)
                    run = p.add_run(strip_markdown_formatting(cell_text))
                    run.font.size = Pt(9)
                    run.font.name = FONT_NAME
                    run.font.color.rgb = colors["text_color"]
                    _set_cell_border(cell, border_color)

                    # 짝수 행 배경색
                    if row_idx % 2 == 0:
                        _set_cell_shading(cell, colors["table_row_even_bg"])

            # 테이블 전후 간격
            p_after = doc.add_paragraph()
            _set_paragraph_spacing(p_after, before=0, after=4)

        elif elem_type == 'code':
            code_text = elem['text']
            # 코드 블록: 배경색 있는 테이블로 표현
            code_table = doc.add_table(rows=1, cols=1)
            code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = code_table.rows[0].cells[0]
            cell.text = ""
            _set_cell_shading(cell, colors["code_bg"])
            _set_cell_border(cell, colors.get("table_border_color", "E2E8F0"))

            for line_idx, code_line in enumerate(code_text.split('\n')):
                if line_idx > 0:
                    p = cell.add_paragraph()
                else:
                    p = cell.paragraphs[0]
                _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)
                run = p.add_run(code_line if code_line else " ")
                run.font.name = CODE_FONT_NAME
                run.font.size = Pt(9)
                run.font.color.rgb = colors["code_text"]

            # 코드 블록 후 간격
            p_after = doc.add_paragraph()
            _set_paragraph_spacing(p_after, before=0, after=4)

        elif elem_type == 'blockquote':
            # 블록인용: 들여쓰기 + 왼쪽 테두리
            bq_text = elem['text']
            for bq_line in bq_text.split('\n'):
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=1)
                p.paragraph_format.left_indent = Cm(1.0)

                # 왼쪽 테두리
                pPr = p._p.get_or_add_pPr()
                bq_border_color = colors.get("blockquote_border", "A0AEC0")
                borders = parse_xml(
                    f'<w:pBorders {nsdecls("w")}>'
                    f'  <w:left w:val="single" w:sz="16" w:space="8" w:color="{bq_border_color}"/>'
                    f'</w:pBorders>'
                )
                pPr.append(borders)

                # 배경색 (shading)
                bg_color = colors.get("blockquote_bg", "F7FAFC")
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
                pPr.append(shd)

                if bq_line.strip():
                    _add_inline_formatted_text(p, bq_line, colors, base_size=10)
                else:
                    run = p.add_run(" ")
                    run.font.size = Pt(10)

        elif elem_type == 'hr':
            _add_horizontal_rule(doc, colors)

        elif elem_type == 'image':
            # 이미지 (차트 등)
            img_path = elem['path']
            caption = elem.get('caption', '')

            # 이미지 파일 찾기
            img_file = _find_image(img_path)
            if img_file:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_file), width=Inches(5.5))

                    if caption:
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_paragraph_spacing(p_cap, before=2, after=6)
                        run_cap = p_cap.add_run(caption)
                        run_cap.font.size = Pt(9)
                        run_cap.font.name = FONT_NAME
                        run_cap.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                        run_cap.italic = True
                except Exception as e:
                    p = doc.add_paragraph()
                    run = p.add_run(f"[이미지 삽입 실패: {e}]")
                    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x66)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[이미지를 찾을 수 없음: {img_path}]")
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return doc


def _find_image(img_path: str) -> Optional[Path]:
    """이미지 파일 찾기"""
    img_file = Path(img_path)
    filename = Path(img_path).name

    if img_file.is_absolute() and img_file.exists():
        return img_file

    possible_bases = [
        Path(r"C:\Users\Administrator\Documents\LFChatbot_NextJS_FastAPI\backend\data\chart_output"),
        Path(__file__).parent.parent.parent.parent / "data" / "chart_output",
        Path.cwd() / "data" / "chart_output",
        Path.cwd() / "backend" / "data" / "chart_output",
    ]

    for base_dir in possible_bases:
        full_path = base_dir / img_path
        if full_path.exists():
            return full_path
        name_path = base_dir / filename
        if name_path.exists():
            return name_path

    return None


# ============================================================================
# MCP 도구 정의
# ============================================================================

@server.list_tools()
async def list_tools():
    return [
        Tool(
            name="create_document_docx",
            description="""마크다운/텍스트 문서를 Word(DOCX)로 변환합니다.

기능:
- 마크다운 문법 지원 (제목 #, 표, 코드블록 ```, 구분선 ---)
- **굵게**, *기울임*, `코드` 인라인 서식 지원
- 불릿 리스트 (-, *, +) 및 번호 리스트 (1. 2. 3.) 지원 (중첩 가능)
- 블록인용 (>) 지원
- 한글 완벽 지원 (맑은 고딕)
- 테이블 자동 스타일링 (컬러 헤더, 줄무늬 배경)
- 코드블록 배경 스타일
- 생성된 문서는 수정 가능 (PDF와 달리 편집 가능)

스타일 옵션:
- technical: 기술 문서용 (파란색 헤더, 깔끔한 표) - 기본값
- report: 보고서용 (공식적인 회색톤)
- simple: 심플한 스타일""",
            inputSchema={
                "type": "object",
                "properties": {
                    "content": {
                        "type": "string",
                        "description": "변환할 마크다운/텍스트 내용"
                    },
                    "title": {
                        "type": "string",
                        "description": "문서 제목"
                    },
                    "subtitle": {
                        "type": "string",
                        "description": "부제목 (선택사항, 제목 아래 이탤릭으로 표시)",
                        "default": ""
                    },
                    "filename": {
                        "type": "string",
                        "description": "출력 파일명 (확장자 제외)"
                    },
                    "style": {
                        "type": "string",
                        "enum": ["technical", "report", "simple"],
                        "default": "technical",
                        "description": "문서 스타일"
                    },
                    "section_per_page": {
                        "type": "boolean",
                        "default": False,
                        "description": "True면 주요 섹션(##)마다 새 페이지에서 시작"
                    }
                },
                "required": ["content", "title", "filename"]
            }
        ),
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict):
    if name == "create_document_docx":
        content = arguments.get("content", "")
        title = arguments.get("title", "문서")
        subtitle = arguments.get("subtitle", "")
        filename = arguments.get("filename", "output")
        style = arguments.get("style", "technical")
        section_per_page = arguments.get("section_per_page", False)

        try:
            safe_filename = re.sub(r'[<>:"/\\|?*`\'"\s]', '', filename).strip()
            output_path = OUTPUT_DIR / f"{safe_filename}.docx"

            # 제목 추가
            content_stripped = content.strip()
            if content_stripped.startswith('# '):
                full_content = content
            else:
                full_content = f"# {title}\n\n{content}"

            elements = parse_markdown_content(full_content)
            doc = render_docx(elements, title, subtitle, style, section_per_page)
            doc.save(str(output_path))

            file_size = output_path.stat().st_size / 1024

            section_hint = ""
            if not section_per_page:
                section_hint = "\n\n[참고] 섹션별로 페이지를 나눠서 다시 생성하려면 section_per_page=true 옵션을 사용하세요."

            return [TextContent(
                type="text",
                text=f"Word 문서 생성 완료\n\n파일: {output_path}\n스타일: {style}\n제목: {title}\n크기: {file_size:.1f} KB{section_hint}"
            )]

        except Exception as e:
            return [TextContent(
                type="text",
                text=f"Word 문서 생성 실패: {str(e)}"
            )]

    return [TextContent(type="text", text=f"알 수 없는 도구: {name}")]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
